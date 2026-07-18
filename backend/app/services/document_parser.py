"""Document parser service.

Handles parsing of PDF, DOCX, TXT, and Markdown files.
Extracts text content and metadata from documents.
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import markdown
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.config import get_settings
from app.core.exceptions import FileUploadException
from app.utils.logging import get_logger

logger = get_logger(__name__)
settings = get_settings()


@dataclass
class ParsedDocument:
    """Parsed document content with metadata."""

    content: str
    metadata: dict = field(default_factory=dict)
    page_count: int = 0
    word_count: int = 0
    char_count: int = 0

    def __post_init__(self):
        """Calculate counts after initialization."""
        if self.content:
            self.char_count = len(self.content)
            self.word_count = len(self.content.split())
            if not self.page_count:
                self.page_count = 1


class DocumentParser:
    """Service for parsing different document formats."""

    def __init__(self):
        """Initialize the document parser."""
        self._parsers = {
            "pdf": self._parse_pdf,
            "docx": self._parse_docx,
            "txt": self._parse_text,
            "md": self._parse_markdown,
            "markdown": self._parse_markdown,
        }

    def parse(self, file_path: str, file_type: str) -> ParsedDocument:
        """Parse a document file and extract text content.

        Args:
            file_path: Path to the document file.
            file_type: File type extension (pdf, docx, txt, md).

        Returns:
            Parsed document with content and metadata.

        Raises:
            FileUploadException: If parsing fails.
        """
        file_type = file_type.lower().strip(".")
        parser = self._parsers.get(file_type)

        if parser is None:
            raise FileUploadException(
                f"Unsupported file type: {file_type}. "
                f"Supported types: {', '.join(self._parsers.keys())}"
            )

        try:
            result = parser(file_path)
            logger.info(
                f"Parsed {file_type} document: {file_path}, "
                f"words={result.word_count}, pages={result.page_count}"
            )
            return result
        except FileUploadException:
            raise
        except Exception as e:
            raise FileUploadException(f"Failed to parse {file_type} file: {str(e)}")

    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """Parse PDF file using pypdf.

        Args:
            file_path: Path to PDF file.

        Returns:
            Parsed document with content and metadata.
        """
        reader = PdfReader(file_path)
        text_parts = []
        page_count = len(reader.pages)

        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        content = "\n\n".join(text_parts)

        metadata = {
            "parser": "pypdf",
            "page_count": page_count,
            "title": reader.metadata.title if reader.metadata else None,
            "author": reader.metadata.author if reader.metadata else None,
        }

        return ParsedDocument(
            content=content,
            metadata=metadata,
            page_count=page_count,
        )

    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse DOCX file using python-docx.

        Args:
            file_path: Path to DOCX file.

        Returns:
            Parsed document with content and metadata.
        """
        doc = DocxDocument(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        content = "\n\n".join(text_parts)

        # Extract document properties
        core_props = doc.core_properties
        metadata = {
            "parser": "python-docx",
            "title": core_props.title,
            "author": core_props.author,
            "subject": core_props.subject,
            "paragraph_count": len(doc.paragraphs),
        }

        return ParsedDocument(
            content=content,
            metadata=metadata,
        )

    def _parse_text(self, file_path: str) -> ParsedDocument:
        """Parse plain text file.

        Args:
            file_path: Path to text file.

        Returns:
            Parsed document with content and metadata.
        """
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        line_count = len(content.splitlines())

        metadata = {
            "parser": "plaintext",
            "line_count": line_count,
        }

        return ParsedDocument(
            content=content,
            metadata=metadata,
        )

    def _parse_markdown(self, file_path: str) -> ParsedDocument:
        """Parse Markdown file, extracting plain text content.

        Args:
            file_path: Path to Markdown file.

        Returns:
            Parsed document with content and metadata.
        """
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            raw_content = f.read()

        # Convert markdown to HTML then strip tags for plain text
        html = markdown.markdown(raw_content)
        text = self._strip_html(html)

        # Extract headings for metadata
        headings = re.findall(r"^#+\s+(.+)$", raw_content, re.MULTILINE)

        metadata = {
            "parser": "markdown",
            "heading_count": len(headings),
            "headings": headings[:10],  # First 10 headings
        }

        return ParsedDocument(
            content=text,
            metadata=metadata,
        )

    def _strip_html(self, html: str) -> str:
        """Strip HTML tags from text.

        Args:
            html: HTML content.

        Returns:
            Plain text content.
        """
        clean = re.sub(r"<[^>]+>", "", html)
        clean = re.sub(r"\s+", " ", clean)
        return clean.strip()


# Module-level instance
document_parser = DocumentParser()
